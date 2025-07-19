import streamlit as st
import fitz  # PyMuPDF
import requests
import time
import re
import os 
import json
import io
import plotly.express as px
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document
from fpdf import FPDF
from cryptography.fernet import Fernet
from dotenv import load_dotenv
load_dotenv()

WHITELIST_FILE = "whitelist.json"
HISTORY_FILE = "login_history.json"
 

from dotenv import load_dotenv
import os
import base64
from cryptography.fernet import Fernet

load_dotenv()  # load variables from .env

secret_key_b64 = os.getenv("SECRET_KEY_BASE64")
encrypted_password_b64 = os.getenv("ENCRYPTED_PASSWORD_BASE64")

# Decode base64 strings back to bytes
secret_key = base64.b64decode(secret_key_b64)
encrypted_password = base64.b64decode(encrypted_password_b64)

def decrypt_password():
    fernet = Fernet(secret_key)
    decrypted = fernet.decrypt(encrypted_password).decode()
    return decrypted

DEV_PASSWORD = decrypt_password()



# ======= Initialize session state and load persistent data =======
def load_whitelist():
    if os.path.exists(WHITELIST_FILE):
        with open(WHITELIST_FILE, "r") as f:
            return json.load(f)
    return []

def save_whitelist(whitelist):
    with open(WHITELIST_FILE, "w") as f:
        json.dump(whitelist, f)

def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r") as f:
            return json.load(f)
    return {}

def save_history(history):
    with open(HISTORY_FILE, "w") as f:
        json.dump(history, f)

if "whitelist" not in st.session_state:
    st.session_state.whitelist = load_whitelist()

if "login_history" not in st.session_state:
    st.session_state.login_history = load_history()

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if "is_dev" not in st.session_state:
    st.session_state.is_dev = False

if "start_time" not in st.session_state:
    st.session_state.start_time = time.time()

# ======= AUTHENTICATION SECTION =======
if not st.session_state.authenticated:
    st.title("🔒 Private version, Login Required")
    choice = st.radio("Login as:", ["Dev", "Guest"])

    if choice == "Dev":
        dev_pass = st.text_input("Enter Dev Password", type="password")
        if st.button("Login as Dev"):
            if dev_pass == DEV_PASSWORD:
                st.session_state.authenticated = True
                st.session_state.is_dev = True
                st.rerun()
            else:
                st.error("Incorrect Dev Password")

    else:  # Friend login
        email = st.text_input("Enter your email")
        if st.button("Login as Guest"):
            if email in st.session_state.whitelist:
                st.session_state.authenticated = True
                st.session_state.is_dev = False
                st.session_state.user_email = email
                # Log history
                now = datetime.now().isoformat()
                hist = st.session_state.login_history
                if email not in hist:
                    hist[email] = []
                hist[email].append({"login": now})
                save_history(hist)
                st.rerun()
            else:
                st.error("Access denied: Email not approved")
    st.stop()





# ======= DEV PANEL (in sidebar) =======
if st.session_state.is_dev:
    with st.sidebar.expander("⚙️ Dev Settings"):
        if st.text_input("Re-enter Dev Password", type="password") == DEV_PASSWORD:
            st.success(f"Welcome Aurumz!")

            st.subheader("Manage Whitelisted Emails")
            new_email = st.text_input("Add email to whitelist")
            if st.button("Add Email") and new_email:
                if new_email not in st.session_state.whitelist:
                    st.session_state.whitelist.append(new_email)
                    save_whitelist(st.session_state.whitelist)
                    st.success(f"{new_email} added")

            st.write("### Current Whitelisted Emails")
            for email in st.session_state.whitelist:
                cols = st.columns([4,1])
                cols[0].write(email)
                if cols[1].button("❌", key=f"del-{email}"):
                    st.session_state.whitelist.remove(email)
                    save_whitelist(st.session_state.whitelist)
                    st.warning(f"{email} removed")

            st.subheader("Login History")
            hist = load_history()
            for email, entries in hist.items():
                st.write(f"**{email}**")
                for record in entries:
                    st.write(f"- {record['login']}")



# Session state initialization
if "note_acknowledged" not in st.session_state:
    st.session_state.note_acknowledged = False

# Show the note at first
if not st.session_state.note_acknowledged:
    st.markdown("""⚠️ **Note:**

 Accuracy of this tool is around 85-90%, with more errors in extracting Secondary Outcomes and non-phrasing issues.

 Upload up to 15 papers maximum, or it may crash and fail to show results.

 Do not re-screen the paper as it will consume 2x tokens.

 $1 ~ 250 Articles, but this varies with article length and the number of Articles can be more or less.

 Always crosscheck especially for maybe/excluded extracted studies.

 Confidence scores are nothing but scores which determine how well texts were readable by this tool (max:0.9).
""")
    if st.button("OK, Got it"):
        st.session_state.note_acknowledged = True
    st.stop()





st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@700&display=swap');

    .stApp {
        background-color: #2B3445;
        color: #F0F4F8;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }

    .typewriter {
        text-align: center;
        font-family: 'Montserrat', sans-serif;
        font-weight: 700;
        font-size: 3.5rem;
        color: #F0F4F8;
    }

    .typewriter .typing {
        display: inline-block;
        overflow: hidden;
        border-right: .15em solid #F0F4F8;
        white-space: nowrap;
        letter-spacing: .04em;
        animation:
            typing 1s steps(11, end) forwards,
            blink-caret 1s step-end forwards;
        vertical-align: bottom;
    }

    .gold {
        color: #FF9B45;
    }

    .typewriter .dot {
        display: inline-block;
        vertical-align: bottom;
        margin-left: 2px;
        color: #F0F4F8;
    }

    @keyframes typing {
        from { width: 0 }
        to { width: 11ch; }
    }

    @keyframes blink-caret {
        0%, 100% { border-color: transparent; }
        50% { border-color: #F0F4F8; }
    }
    </style>

    <div class="typewriter">
        <span class="typing">Review<span class="gold">Aid</span>.</span>
    </div>

    <h3 style='text-align: center; color: #F0F4F8; margin-top: 10px;'>
        A Research Article Screener & Extractor
    </h3>
    """,
    unsafe_allow_html=True
)





# ========== CONFIGURATION ==========
from dotenv import load_dotenv
import os

load_dotenv()  # loads environment variables from .env file

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

if DEEPSEEK_API_KEY is None:
    raise ValueError("API key not found. Please set it in the .env file")

if "included_results" not in st.session_state:
    st.session_state.included_results = []
if "excluded_results" not in st.session_state:
    st.session_state.excluded_results = []

if "maybe_results" not in st.session_state:
    st.session_state.maybe_results = []
if "start_time" not in st.session_state:
    st.session_state.start_time = time.time()
if "decision_timestamps" not in st.session_state:
    st.session_state.decision_timestamps = []


# ========== FUNCTIONS ==========

def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    return full_text


def query_deepseek(prompt, api_key=DEEPSEEK_API_KEY):
    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3
    }
    response = requests.post(url, json=payload, headers=headers)
    response.raise_for_status()
    return response.json()['choices'][0]['message']['content']

import json
import re
import streamlit as st

def parse_result(result_str):
    try:
        match = re.search(r"\{.*\}", result_str, re.DOTALL)
        if match:
            json_str = match.group(0)
            data = json.loads(json_str)
            # Confirm keys exist
            if "extracted" not in data:
                st.warning("Extracted fields missing in AI response")
            return data
        else:
            st.error("No JSON object found in API response.")
            return None
    except Exception as e:
        st.error(f"JSON parsing error: {e}")
        return None
    

def df_from_results(results):
    rows = []
    for r in results:
        row = {
    "Filename": r.get("filename", ""),
    "Status": r.get("status", "").capitalize(),
    "Confidence": r.get("confidence", "")
}
        row.update(r.get("extracted", {}))
        if r.get("status", "").lower() == "exclude":
            row["Reason for Exclusion"] = r.get("reason", "")
        rows.append(row)
    return pd.DataFrame(rows)

def estimate_confidence(text):
    if not text or len(text.strip()) < 30:
        return 0.2
    if "randomized" in text.lower():
        return 0.9
    return 0.6


def to_docx(df):
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    doc.add_heading('Exported Papers', 0)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    table.autofit = False

    # Set header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        hdr_cells[i].width = Inches(2)  # Optional fixed width

    # Add data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            text = str(val)
            para = row_cells[i].paragraphs[0]
            para.add_run(text)

            # Enable word wrap (in case)
            tc = row_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(OxmlElement('w:vAlign'))
            tcPr[-1].set(qn('w:val'), 'top')

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def to_pdf(df):
    from fpdf import FPDF

    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.set_font("Arial", size=8)

    col_width = 280 / len(df.columns)
    row_height = 6
    line_height = 4

    # Headers
    for col in df.columns:
        pdf.multi_cell(col_width, line_height, str(col), border=1, align='C')
    pdf.ln(row_height)

    # Rows
    for _, row in df.iterrows():
        y_before = pdf.get_y()
        max_y = y_before
        x_start = pdf.get_x()

        for i, val in enumerate(row):
            text = str(val)
            x = pdf.get_x()
            y = pdf.get_y()
            pdf.multi_cell(col_width, line_height, text, border=1)
            pdf.set_xy(x + col_width, y)
            max_y = max(max_y, pdf.get_y())
        pdf.set_y(max_y)

    pdf_str = pdf.output(dest='S')  # get PDF as string
    pdf_bytes = pdf_str.encode('latin1')  # encode string to bytes

    return pdf_bytes



def to_csv(df):
    return df.to_csv(index=False).encode('utf-8')

def to_excel(df):
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')
    buffer.seek(0)
    return buffer.getvalue()


# ========== STREAMLIT APP ==========



st.subheader("Population Criteria")
population_inclusion = st.text_area("Population Inclusion Criteria", placeholder="e.g. Adults aged 18–65 with MS")
population_exclusion = st.text_area("Population Exclusion Criteria", placeholder="e.g. Patients with comorbid autoimmune diseases")

st.subheader("Intervention Criteria")
intervention_inclusion = st.text_area("Intervention Inclusion Criteria", placeholder="e.g. Natalizumab treatment ≥ 6 months")
intervention_exclusion = st.text_area("Intervention Exclusion Criteria", placeholder="e.g. Dose outside approved range")

st.subheader("Comparison Criteria")
comparison_inclusion = st.text_area("Comparison Inclusion Criteria", placeholder="e.g. Placebo or no treatment")
comparison_exclusion = st.text_area("Comparison Exclusion Criteria", placeholder="e.g. Active comparator like interferon beta")

st.subheader("Outcome Criteria (Optional)")
outcome_criteria = st.text_area("Outcome Criteria", placeholder="e.g. Annualized relapse rate, disability progression")

fields = st.text_input("Fields to Extract (comma-separated)", placeholder="e.g. Author, Year,Population, Outcome")
uploaded_pdfs = st.file_uploader("Upload PDF Files", accept_multiple_files=True)

if st.button("Screen & Extract"):

    if not uploaded_pdfs:
        st.warning("Please upload at least one PDF file.")
        st.stop()

    if not any([
       population_inclusion.strip(), population_exclusion.strip(),
       intervention_inclusion.strip(), intervention_exclusion.strip(),
       comparison_inclusion.strip(), comparison_exclusion.strip(),
       outcome_criteria.strip()
    ]):
       st.warning("Please enter at least one inclusion or exclusion criterion.")
       st.stop()

    if not fields.strip():
        st.warning("Please specify at least one field to extract.")
        st.stop()

    fields_list = [f.strip() for f in fields.split(",") if f.strip()]

    included_results = []
    excluded_results = []
    maybe_results = []

    for pdf in uploaded_pdfs:
       with st.spinner(f"Reading {pdf.name}..."):
          text = extract_text_from_pdf(pdf)

       confidence = estimate_confidence(text)
       if confidence < 0.5:
          st.warning(f"Low Confidence in abstract: {confidence:.2f}")
       else:
          st.success(f"High Confidence: {confidence:.2f}")

       prompt = f"""
You are screening a research paper.

Apply the following structured inclusion and exclusion criteria:

**Population**
Inclusion: {population_inclusion}
Exclusion: {population_exclusion}

**Intervention**
Inclusion: {intervention_inclusion}
Exclusion: {intervention_exclusion}

**Comparison**
Inclusion: {comparison_inclusion}
Exclusion: {comparison_exclusion}

**Outcomes (if relevant)**: {outcome_criteria}

Fields to extract: {', '.join(fields_list)}

Here is the full text of the paper:
{text}

Return this as JSON:
{{
  "status": "Include/Exclude/Maybe",
  "reason": "Reason for classification",
  "extracted": {{
  {', '.join(f'"{field}": "..."' for field in fields_list)}
  }}
}}
"""

       
       with st.spinner(f"Sending '{pdf.name}' to DeepSeek AI..."):
        try:
            raw_result = query_deepseek(prompt)
            result = parse_result(raw_result)
            if not result:
                st.error(f"Could not parse result for {pdf.name}. Raw output:")
                st.code(raw_result)
                continue  # continue here is valid because inside the for loop

            result["filename"] = pdf.name
            result["confidence"] = confidence
            if confidence < 0.5:
                result["flags"] = ["low_confidence"]

            status = result.get("status", "").lower()
            if status == "include":
                st.session_state.included_results.append(result)
                st.session_state.decision_timestamps.append(time.time())
            elif status == "exclude":
                st.session_state.excluded_results.append(result)
                st.session_state.decision_timestamps.append(time.time())
            elif status == "maybe":
                st.session_state.maybe_results.append(result)
                st.session_state.decision_timestamps.append(time.time())
            else:
                st.session_state.excluded_results.append(result)
                st.session_state.decision_timestamps.append(time.time())

            st.success(f"Processed: {pdf.name} — {status.capitalize()}")
            
        except Exception as e:
               st.error(f"Error processing {pdf.name}: {str(e)}")




import plotly.express as px

included = len(st.session_state.included_results)
excluded = len(st.session_state.excluded_results)
maybe = len(st.session_state.maybe_results)
total = included + excluded + maybe

session_duration = time.time() - st.session_state.start_time
avg_speed = session_duration / total if total > 0 else 0

with st.expander("Screening Dashboard", expanded=True):
     st.metric("Papers Screened", total)
     st.metric("Session Time (min) ", f"{session_duration / 60 :.2f}")
     st.metric("Avg Decision Time (s)", f"{avg_speed:.2f}")

     fig = px.pie(
        names=["Included", "Excluded", "Maybe"],
        values=[included, excluded, maybe],
        title="Screening Decisions"
    )
     st.plotly_chart(fig, use_container_width=True)

                

    # Show summary tables
included_results = st.session_state.included_results
excluded_results = st.session_state.excluded_results
maybe_results = st.session_state.maybe_results 

if included_results:
    st.header("Included Papers")
    df_inc = df_from_results(included_results)
    st.dataframe(df_inc)
else:
    st.info("No Included papers found.")

if excluded_results:
    st.header("Excluded Papers")
    df_exc = df_from_results(excluded_results)
    st.dataframe(df_exc)
else:
    st.info("No Excluded papers found.")

    maybe_results = st.session_state.maybe_results

if maybe_results:
    st.header("Maybe Papers")
    df_maybe = df_from_results(maybe_results)
    st.dataframe(df_maybe)
else:
    st.info("No Maybe papers found.")

# Export options # Export options (DOCX, CSV, XLSX)
if included_results or excluded_results or maybe_results:
    st.header("Export Results")

    def export_buttons(df, label_prefix):
        formats = {
            "DOCX": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", to_docx(df), f"{label_prefix.lower()}_papers.docx"),
            "CSV":  ("text/csv", to_csv(df), f"{label_prefix.lower()}_papers.csv"),
            "XLSX": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", to_excel(df), f"{label_prefix.lower()}_papers.xlsx")
        }

        for fmt, (mime, data, filename) in formats.items():
            st.download_button(
                label=f"Download {label_prefix} as {fmt}",
                data=data,
                file_name=filename,
                mime=mime
            )

    if included_results:
        st.subheader("Included Papers")
        df_inc = df_from_results(included_results)
        export_buttons(df_inc, "Included")

    if excluded_results:
        st.subheader("Excluded Papers")
        df_exc = df_from_results(excluded_results)
        export_buttons(df_exc, "Excluded")

    if maybe_results:
        st.subheader("Maybe Papers")
        df_maybe = df_from_results(maybe_results)
        export_buttons(df_maybe, "Maybe")



import subprocess
import streamlit as st

# Get version from latest tag
try:
    version = subprocess.check_output(["git", "describe", "--tags"]).decode().strip()
except:
    version = "v?.?.?"  # fallback

# Get last updated date
try:
    last_updated = subprocess.check_output(
        ["git", "log", "-1", "--format=%cd", "--date=short"]
    ).decode().strip()
except:
    last_updated = "Unknown"

# Display footer
st.markdown(
    f"""
    <style>
    .custom-footer-container {{
        width: 100%;
        font-family: 'Times New Roman', Times, serif;
        font-size: 10px;
        color: #F0F4F8;
        opacity: 0.7;
        margin-top: 150px;
        padding: 0 10px;
        overflow: hidden;
        position: fixed;
        bottom: 0;
        left: 0;
    }}

    .footer-left {{
        float: left;
    }}

    .footer-right {{
        float: right;
    }}
    </style>

    <div class="custom-footer-container">
        <div class="footer-left">
            Version {version} &nbsp;|&nbsp; Last updated: {last_updated}
        </div>
        <div class="footer-right">
            Made with 💛 by its Creator.
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

